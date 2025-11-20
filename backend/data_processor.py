"""
Enhanced Data Processor Module with Comprehensive OCR Support
============================================================
This module handles document processing and text extraction with advanced OCR:
- PDF text extraction with embedded image OCR
- DOCX/DOC text extraction with table support
- Image OCR processing (multi-language support)
- Excel/CSV table extraction
- Google Gemini Vision API integration for advanced OCR
- Fallback to Tesseract OCR and EasyOCR

وحدة معالج البيانات المحسّنة مع دعم OCR شامل
==============================================
هذه الوحدة تتعامل مع معالجة المستندات واستخراج النص مع OCR متقدم:
- استخراج النص من PDF مع OCR للصور المضمنة
- استخراج النص من DOCX/DOC مع دعم الجداول
- معالجة OCR للصور (دعم متعدد اللغات)
- استخراج جداول Excel/CSV
- تكامل Google Gemini Vision API لـ OCR متقدم
- العودة إلى Tesseract OCR و EasyOCR
"""

import os
import logging
import base64
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

# Core libraries
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pandas as pd

logger = logging.getLogger("DATA_PROCESSOR")

# OCR libraries
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    EASYOCR_READER = None  # Will be initialized on first use
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False

from langchain_core.documents import Document

logger = logging.getLogger("DATA_PROCESSOR")

# ------------------------------------------------------------
# Configuration
# ------------------------------------------------------------
GEMINI_API_KEY = os.getenv("GOOGLE_GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_VISION_MODEL", "gemini-1.5-pro-vision")
USE_GEMINI_OCR = os.getenv("USE_GEMINI_OCR", "true").lower() == "true"
OCR_LANGUAGES = ['ara', 'eng']  # Arabic and English
MAX_IMAGE_SIZE = 20 * 1024 * 1024  # 20MB max for Gemini API

# Initialize Gemini if available
if GEMINI_AVAILABLE and GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        logger.info("Google Gemini API configured successfully")
    except Exception as e:
        logger.warning(f"Failed to configure Gemini API: {e}")
        GEMINI_AVAILABLE = False

# Initialize EasyOCR reader on first use
def _get_easyocr_reader():
    """Lazy initialization of EasyOCR reader."""
    global EASYOCR_READER
    if EASYOCR_AVAILABLE and EASYOCR_READER is None:
        try:
            EASYOCR_READER = easyocr.Reader(['ar', 'en'], gpu=False)
            logger.info("EasyOCR reader initialized successfully")
        except Exception as e:
            logger.warning(f"Failed to initialize EasyOCR: {e}")
    return EASYOCR_READER


# ------------------------------------------------------------
# Helper Functions
# ------------------------------------------------------------

def _encode_image_to_base64(image_path: str) -> Optional[str]:
    """
    Encode image to base64 for API calls.
    / ترميز الصورة إلى base64 لاستدعاءات API.
    """
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except Exception as e:
        logger.error(f"Error encoding image {image_path}: {e}")
        return None


def _get_image_size(image_path: str) -> int:
    """Get image file size in bytes."""
    try:
        return os.path.getsize(image_path)
    except Exception:
        return 0


def _resize_image_if_needed(image_path: str, max_size: int = MAX_IMAGE_SIZE) -> str:
    """
    Resize image if it exceeds max size.
    / تغيير حجم الصورة إذا تجاوزت الحجم الأقصى.
    """
    try:
        if _get_image_size(image_path) <= max_size:
            return image_path
        
        img = Image.open(image_path)
        # Calculate new dimensions maintaining aspect ratio
        ratio = (max_size / _get_image_size(image_path)) ** 0.5
        new_size = (int(img.width * ratio), int(img.height * ratio))
        img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
        
        # Save to temporary file
        temp_path = image_path + ".resized"
        img_resized.save(temp_path, format=img.format)
        logger.info(f"Resized image from {_get_image_size(image_path)} to {_get_image_size(temp_path)} bytes")
        return temp_path
    except Exception as e:
        logger.warning(f"Failed to resize image: {e}")
        return image_path


# ------------------------------------------------------------
# OCR Functions
# ------------------------------------------------------------

def _ocr_with_gemini(image_path: str) -> Optional[str]:
    """
    Extract text from image using Google Gemini Vision API.
    / استخراج النص من الصورة باستخدام Google Gemini Vision API.
    
    Args:
        image_path: Path to image file / مسار ملف الصورة
        
    Returns:
        Extracted text or None / النص المستخرج أو None
    """
    if not GEMINI_AVAILABLE or not GEMINI_API_KEY or not USE_GEMINI_OCR:
        return None
    
    try:
        # Resize if needed
        processed_path = _resize_image_if_needed(image_path)
        
        # Load and encode image
        img = Image.open(processed_path)
        
        # Use Gemini Vision API
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        prompt = """
        Extract all text from this image. 
        The image may contain:
        - Arabic text (right-to-left)
        - English text (left-to-right)
        - Tables with structured data
        - Mixed content
        
        Please extract:
        1. All visible text maintaining reading order
        2. Table data in a structured format (markdown table if applicable)
        3. Preserve formatting where possible
        
        Respond with the extracted text only, no explanations.
        """
        
        response = model.generate_content([prompt, img])
        text = response.text
        
        # Clean up temporary file if created
        if processed_path != image_path and os.path.exists(processed_path):
            try:
                os.remove(processed_path)
            except Exception:
                pass
        
        logger.info(f"Successfully extracted text using Gemini Vision API from {os.path.basename(image_path)}")
        return text
        
    except Exception as e:
        logger.error(f"Error in Gemini OCR for {image_path}: {e}")
        return None


def _ocr_with_easyocr(image_path: str) -> Optional[str]:
    """
    Extract text from image using EasyOCR.
    / استخراج النص من الصورة باستخدام EasyOCR.
    """
    if not EASYOCR_AVAILABLE:
        return None
    
    try:
        reader = _get_easyocr_reader()
        if reader is None:
            return None
        
        results = reader.readtext(image_path)
        text = "\n".join([result[1] for result in results])
        logger.info(f"Successfully extracted text using EasyOCR from {os.path.basename(image_path)}")
        return text
    except Exception as e:
        logger.error(f"Error in EasyOCR for {image_path}: {e}")
        return None


def _ocr_with_tesseract(image_path: str) -> Optional[str]:
    """
    Extract text from image using Tesseract OCR.
    / استخراج النص من الصورة باستخدام Tesseract OCR.
    """
    if not TESSERACT_AVAILABLE:
        return None
    
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='ara+eng')
        logger.info(f"Successfully extracted text using Tesseract from {os.path.basename(image_path)}")
        return text
    except Exception as e:
        logger.error(f"Error in Tesseract OCR for {image_path}: {e}")
        return None


def _extract_text_from_image_advanced(image_path: str) -> str:
    """
    Extract text from image using multiple OCR methods with fallback.
    / استخراج النص من الصورة باستخدام طرق OCR متعددة مع العودة.
    
    Priority:
    1. Google Gemini Vision API (if available and enabled)
    2. EasyOCR (if available)
    3. Tesseract OCR (if available)
    """
    text = None
    
    # Try Gemini first (best quality)
    if USE_GEMINI_OCR:
        text = _ocr_with_gemini(image_path)
        if text and text.strip():
            return text
    
    # Fallback to EasyOCR
    if not text or not text.strip():
        text = _ocr_with_easyocr(image_path)
        if text and text.strip():
            return text
    
    # Final fallback to Tesseract
    if not text or not text.strip():
        text = _ocr_with_tesseract(image_path)
        if text and text.strip():
            return text
    
    if not text or not text.strip():
        logger.warning(f"Failed to extract text from image {image_path} using all OCR methods")
        return ""
    
    return text


# ------------------------------------------------------------
# Document Extraction Functions
# ------------------------------------------------------------

def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file with comprehensive OCR support.
    / استخراج النص من ملف PDF مع دعم OCR شامل.
    
    This function:
    1. Extracts text directly from PDF
    2. Extracts tables from PDF
    3. Performs OCR on pages that are images
    4. Handles embedded images
    """
    full_text = ""
    tables_text = ""
    
    try:
        # Method 1: Extract text directly using pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                logger.debug(f"Processing PDF page {page_num}/{len(pdf.pages)}")
                
                # Extract regular text
                text = page.extract_text()
                if text:
                    full_text += f"\n\n--- Page {page_num} ---\n\n{text}\n"
                
                # Extract tables using pdfplumber
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            # Convert table to markdown format
                            table_md = _table_to_markdown(table)
                            tables_text += f"\n\n--- Table {table_num} (Page {page_num}) ---\n{table_md}\n"
                
                # If page has no text, it might be an image - try OCR
                if not text or len(text.strip()) < 50:
                    if PDF2IMAGE_AVAILABLE:
                        try:
                            # Convert PDF page to image
                            images = convert_from_path(file_path, first_page=page_num, last_page=page_num, dpi=300)
                            if images:
                                temp_image_path = f"/tmp/pdf_page_{page_num}.png"
                                images[0].save(temp_image_path, "PNG")
                                
                                # Perform OCR on the image
                                ocr_text = _extract_text_from_image_advanced(temp_image_path)
                                if ocr_text:
                                    full_text += f"\n\n--- OCR Text from Page {page_num} ---\n{ocr_text}\n"
                                
                                # Clean up
                                try:
                                    os.remove(temp_image_path)
                                except Exception:
                                    pass
                        except Exception as e:
                            logger.debug(f"Could not perform OCR on page {page_num}: {e}")
        
        # Combine text and tables
        result = full_text + tables_text
        return result
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
        return full_text


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file including tables.
    / استخراج النص من ملف DOCX بما في ذلك الجداول.
    """
    full_text = ""
    try:
        doc = DocxDocument(file_path)
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
        
        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                table_md = _table_to_markdown(table_data)
                full_text += f"\n\n--- Table {table_num} ---\n{table_md}\n"
        
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
        return ""


def _extract_text_from_image(file_path: str) -> str:
    """
    Extract text from image file using advanced OCR.
    / استخراج النص من ملف صورة باستخدام OCR متقدم.
    """
    return _extract_text_from_image_advanced(file_path)


def _extract_text_from_excel(file_path: str) -> str:
    """
    Extract text from Excel file (XLSX, XLS).
    / استخراج النص من ملف Excel (XLSX, XLS).
    """
    full_text = ""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            full_text += f"\n\n--- Sheet: {sheet_name} ---\n"
            full_text += df.to_markdown(index=False) + "\n"
        
        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from Excel {file_path}: {e}", exc_info=True)
        return ""


def _extract_text_from_csv(file_path: str) -> str:
    """
    Extract text from CSV file.
    / استخراج النص من ملف CSV.
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1256', 'windows-1256']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if df is None:
            df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
        
        return df.to_markdown(index=False)
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path}: {e}", exc_info=True)
        return ""


def _extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file with multiple encoding support.
    / استخراج النص من ملف TXT مع دعم ترميزات متعددة.
    """
    try:
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1256', 'windows-1256']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Final fallback
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}", exc_info=True)
        return ""


def _table_to_markdown(table: List[List[str]]) -> str:
    """
    Convert table data to markdown format.
    / تحويل بيانات الجدول إلى صيغة markdown.
    """
    if not table or not table[0]:
        return ""
    
    # Use first row as header
    header = table[0]
    rows = table[1:] if len(table) > 1 else []
    
    # Create markdown table
    md = "| " + " | ".join(str(cell) for cell in header) + " |\n"
    md += "| " + " | ".join("---" for _ in header) + " |\n"
    
    for row in rows:
        # Pad row if needed
        padded_row = row + [""] * (len(header) - len(row))
        md += "| " + " | ".join(str(cell) for cell in padded_row[:len(header)]) + " |\n"
    
    return md


# ------------------------------------------------------------
# Main Processing Functions
# ------------------------------------------------------------

def process_document(file_path: str) -> Optional[Document]:
    """
    Process a single document and extract text to create a Document object.
    / معالجة ملف واحد واستخراج النص لإنشاء كائن Document.
    
    Supports:
    - PDF (with OCR for image-based PDFs)
    - DOCX, DOC (with table extraction)
    - TXT (multiple encodings)
    - Images: JPG, JPEG, PNG, TIFF, BMP, WEBP (with advanced OCR)
    - Excel: XLSX, XLS
    - CSV
    
    Args:
        file_path: Path to the document file / مسار ملف المستند
        
    Returns:
        Document object with extracted text and metadata / كائن Document مع النص المستخرج والبيانات الوصفية
    """
    filename = os.path.basename(file_path)
    file_ext = Path(filename).suffix.lower()
    full_text = ""
    
    logger.info(f"Processing document: {filename} (type: {file_ext})")
    
    try:
        # Route to appropriate extractor
        if file_ext == ".pdf":
            full_text = _extract_text_from_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            full_text = _extract_text_from_docx(file_path)
        elif file_ext == ".txt":
            full_text = _extract_text_from_txt(file_path)
        elif file_ext in [".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp"]:
            full_text = _extract_text_from_image(file_path)
        elif file_ext in [".xlsx", ".xls"]:
            full_text = _extract_text_from_excel(file_path)
        elif file_ext == ".csv":
            full_text = _extract_text_from_csv(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext} for file {filename}")
            return None
        
        # Validate extracted text
        if not full_text or not full_text.strip():
            logger.warning(f"No text extracted from {filename}")
            return None
        
        # Create Document object with metadata
        metadata = {
            "source": filename,
            "file_type": file_ext,
            "content_length": len(full_text),
            "processing_method": "enhanced_ocr" if file_ext in [".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp", ".pdf"] else "standard"
        }
        
        logger.info(f"Successfully processed {filename}: {len(full_text)} characters extracted")
        return Document(page_content=full_text, metadata=metadata)
        
    except Exception as e:
        logger.error(f"Error processing document {filename}: {e}", exc_info=True)
        return None


def ingest_all_documents(data_dir: str) -> List[Document]:
    """
    Index all supported documents in the data directory.
    / فهرسة جميع المستندات المدعومة في مجلد البيانات.
    
    Args:
        data_dir: Path to data directory / مسار مجلد البيانات
        
    Returns:
        List of Document objects / قائمة كائنات Document
    """
    loaded_docs = []
    
    if not os.path.exists(data_dir):
        logger.error(f"Data directory does not exist: {data_dir}")
        return loaded_docs
    
    logger.info(f"Scanning directory: {data_dir}")
    
    # Get all files
    files = [f for f in os.listdir(data_dir) if os.path.isfile(os.path.join(data_dir, f))]
    logger.info(f"Found {len(files)} files in directory")
    
    # Supported extensions
    supported_extensions = {
        ".pdf", ".docx", ".doc", ".txt",
        ".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp",
        ".xlsx", ".xls", ".csv"
    }
    
    for filename in files:
        file_path = os.path.join(data_dir, filename)
        file_ext = Path(filename).suffix.lower()
        
        # Skip unsupported files
        if file_ext not in supported_extensions:
            logger.debug(f"Skipping unsupported file: {filename}")
            continue
        
        logger.info(f"Processing file: {filename}")
        
        try:
            doc = process_document(file_path)
            if doc and doc.page_content and doc.page_content.strip():
                loaded_docs.append(doc)
                logger.info(f"✓ Successfully processed {filename} ({len(doc.page_content)} chars)")
            else:
                logger.warning(f"✗ Failed to extract content from {filename}")
        except Exception as e:
            logger.error(f"✗ Error processing {filename}: {e}", exc_info=True)
    
    logger.info(f"Total documents successfully loaded: {len(loaded_docs)}/{len(files)}")
    return loaded_docs
